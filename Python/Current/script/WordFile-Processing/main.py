import time

from icecream import ic

from docx import Document

doc = Document()


# Adding header for the doc file
doc.add_heading('File Header.', 0)


# Paragraph sample
paragraph = doc.add_paragraph('This is a paragraph sample.')
# Customizing paragraph text
paragraph.add_run(' This is a BOLDED text customization.').bold = True
paragraph.add_run(' This is an ITALIC text customization.').italic = True


doc.add_paragraph('This is item zero', style='List Bullet')
doc.add_paragraph('This is item one', style='List Bullet')
doc.add_paragraph('This is item two', style='List Bullet')
doc.add_paragraph('This is item three', style='List Bullet')
doc.add_paragraph('This is item four', style='List Bullet')


table_header = ['Name', 'Age', 'Greade']

table_data = [
    ['Bruno', 22, 'Ensino Médio Completo'],
    ['Julia', 22, 'Computer Science HighSchool'],
    ['Ana', 23, 'System Developer Technician'],
    ['Nathalia', 38, 'Tech Manager']
]

table = doc.add_table(rows=1, cols=3)
for i in range(3):
    cell = table.rows[0].cells[i]
    cell.paragraphs[0].add_run(table_header[i]).bold = True

for name, age, grade in table_data:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = str(age)
    cells[2].text = grade


# Save file
doc.save('test.docx')

